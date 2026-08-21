#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Genera la planeacion didactica en Word a partir de un archivo JSON.

Respeta el formato oficial del Colegio de Bachilleres (las cinco tablas:
identificacion, elementos curriculares, desarrollos con transversalidad, y una
tabla por cada fase) y adapta las etiquetas al modelo del programa de origen:
UAC (Modelo 2023) o Asignatura (Modelo Educativo 2025).

Uso:
    python generar_planeacion.py planeacion.json salida.docx
    python generar_planeacion.py --esquema        # imprime el JSON de ejemplo
"""

import sys
import os
import json

sys.path.insert(0, os.path.join(
    os.path.dirname(os.path.abspath(__file__)), '..', '..',
    'docx-cobach', 'scripts'))

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    sys.stderr.write('Falta python-docx. Instalar con: pip install python-docx lxml\n')
    sys.exit(2)

try:
    import docx_cobach as dc
except ImportError:
    sys.stderr.write(
        'No se encontro el motor docx-cobach. Se esperaba en\n'
        '  .claude/skills/docx-cobach/scripts/docx_cobach.py\n')
    sys.exit(2)

import terminologia as T

# Se usa el perfil `documento-tecnico`, no `guia`, por dos razones. La primera
# es una regla vigente del Mtro. Luis Gabriel: el perfil `guia` no entra al flujo
# real hasta pasar una prueba de fidelidad contra una guia ya aprobada. La
# segunda es practica: la planeacion son tablas anchas y los margenes de 2.2 cm
# de `documento-tecnico` les dan casi cuatro centimetros mas de caja util que
# los 3.75 cm de `guia`.
PERFIL_NOMBRE = 'documento-tecnico'
PERFIL = dc.PERFILES[PERFIL_NOMBRE]
VERDE = PERFIL['tabla_encabezado_fondo']       # 006837
MENTA = dc.PERFILES['guia']['tabla_cebra']     # D1EBE5, verde menta
FUENTE = PERFIL['fuente']
PT_TABLA = 9.0


# --------------------------------------------------------------------------
# Utilidades de celda
# --------------------------------------------------------------------------

def _texto(celda, texto, negrita=False, blanco=False, tam=PT_TABLA,
           centrado=False):
    celda.text = ''
    p = celda.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if centrado:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, linea in enumerate(str(texto).split('\n')):
        if i:
            p = celda.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            if centrado:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(linea)
        r.font.name = FUENTE
        r.font.size = Pt(tam)
        r.font.bold = negrita
        if blanco:
            r.font.color.rgb = dc.BLANCO
    return celda


def _enc(celda, texto, tam=PT_TABLA):
    """Celda de encabezado: fondo verde, texto blanco, centrado."""
    dc.sombrear_celda(celda, VERDE)
    return _texto(celda, texto, negrita=True, blanco=True, tam=tam, centrado=True)


def _sub(celda, texto, tam=PT_TABLA):
    """Subencabezado: fondo menta, texto negro."""
    dc.sombrear_celda(celda, MENTA)
    return _texto(celda, texto, negrita=True, tam=tam, centrado=True)


def _etq(celda, texto, tam=PT_TABLA):
    """Celda de etiqueta en la columna izquierda."""
    dc.sombrear_celda(celda, PERFIL['ficha_etiqueta_fondo'])
    return _texto(celda, texto, negrita=True, tam=tam)


def _tabla(doc, filas, cols, anchos=None):
    t = doc.add_table(rows=filas, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    dc.borde_tabla_negro(t)
    t.autofit = False
    if anchos:
        for fila in t.rows:
            for i, c in enumerate(fila.cells):
                if i < len(anchos):
                    c.width = Cm(anchos[i])
    return t


def _fusionar_fila(tabla, indice, texto, encabezado=True):
    fila = tabla.rows[indice]
    celda = fila.cells[0]
    for otra in fila.cells[1:]:
        celda = celda.merge(otra)
    if encabezado:
        _enc(celda, texto)
    else:
        _texto(celda, texto)
    return celda


# --------------------------------------------------------------------------
# Bloques del documento
# --------------------------------------------------------------------------

def bloque_identificacion(doc, ident, voc):
    _tabla_id = _tabla(doc, 1, 4, anchos=[4.0, 5.5, 3.2, 3.3])
    _fusionar_fila(_tabla_id, 0, 'Datos generales de identificación')

    pares = [
        ('Institución:', ident.get('institucion', 'Colegio de Bachilleres')),
        ('Nombre del profesor:', ident.get('docente', '')),
        ('Plantel:', ident.get('plantel', '')),
        ('Currículum', ident.get('curriculum', 'Laboral')),
        ('TOB:', ident.get('tob', '')),
        # Aqui cambia la etiqueta segun el modelo: "UAC:" o "Asignatura:"
        (voc['unidad'] + ':', ident.get('unidad', '')),
    ]
    for etq, val in pares:
        fila = _tabla_id.add_row()
        c = fila.cells[0]
        for otra in fila.cells[1:2]:
            c = c.merge(otra)
        _etq(c, etq)
        v = fila.cells[-2].merge(fila.cells[-1])
        _texto(v, val)

    # Dos filas de cuatro columnas
    for a, b, c_, d in [
        ('Semestre:', ident.get('semestre', ''),
         'Ciclo escolar:', ident.get('ciclo', '')),
        ('Horas semestre:', ident.get('horas_semestre', ''),
         'Fecha:', ident.get('fecha', '')),
    ]:
        fila = _tabla_id.add_row()
        _etq(fila.cells[0], a)
        _texto(fila.cells[1], b)
        _etq(fila.cells[2], c_)
        _texto(fila.cells[3], d)
    doc.add_paragraph()


def bloque_curricular(doc, corte, voc):
    t = _tabla(doc, 1, 4, anchos=[4.0, 5.5, 3.2, 3.3])
    _fusionar_fila(t, 0, 'Elementos curriculares')

    def fila_larga(etiqueta, valor):
        fila = t.add_row()
        c = fila.cells[0].merge(fila.cells[1])
        _etq(c, etiqueta)
        v = fila.cells[-2].merge(fila.cells[-1])
        _texto(v, valor)

    fila_larga('Competencia(s) laboral(es):', corte.get('competencia_laboral', ''))
    fila_larga(voc['proposito'] + ':', corte.get('proposito', ''))

    fila = t.add_row()
    _etq(fila.cells[0], 'Número y nombre del corte:')
    _texto(fila.cells[1], corte.get('nombre_corte', ''))
    _etq(fila.cells[2], 'Horas del corte:')
    _texto(fila.cells[3], corte.get('horas_corte', ''))

    fila_larga(voc['meta_corte'] + ':', corte.get('meta_corte', ''))
    doc.add_paragraph()


EVIDENCIAS = (('conocimiento', 'Conocimiento'),
              ('desempeno', 'Desempeño'),
              ('producto', 'Producto'))

ANCHOS_EV = {1: [7.0, 8.6], 2: [6.0, 4.8, 4.8], 3: [6.0, 3.4, 3.1, 3.1]}


def bloque_desarrollos(doc, corte, voc):
    """Un bloque de evidencias por desarrollo, con las columnas que ese
    desarrollo declara y ninguna mas.

    No todos los desarrollos traen las tres evidencias: unos solo declaran
    conocimientos, otros desempeno y producto. Dibujar siempre tres columnas
    dejaba celdas vacias que se leen como un descuido. Ademas es la forma en que
    el propio programa maqueta sus tablas, con un encabezado de Evidencias por
    cada desarrollo.
    """
    desarrollos = corte.get('desarrollos', []) or [{}]

    for d in desarrollos:
        presentes = [(k, etq) for k, etq in EVIDENCIAS if (d.get(k) or '').strip()]
        n = len(presentes)
        anchos = ANCHOS_EV.get(n, ANCHOS_EV[3])
        cols = 1 + max(n, 1)

        t = _tabla(doc, 2, cols, anchos=anchos)
        izq = t.cell(0, 0).merge(t.cell(1, 0))
        _enc(izq, voc['desarrollos'])

        if n:
            der = t.cell(0, 1)
            for j in range(2, cols):
                der = der.merge(t.cell(0, j))
            _enc(der, 'Evidencias de aprendizaje')
            for j, (_, etq) in enumerate(presentes, start=1):
                _sub(t.cell(1, j), etq)
        else:
            der = t.cell(0, 1).merge(t.cell(1, 1))
            _enc(der, 'Evidencias de aprendizaje')

        fila = t.add_row()
        _texto(fila.cells[0], d.get('texto', ''))
        if n:
            for j, (k, _) in enumerate(presentes, start=1):
                _texto(fila.cells[j], d.get(k, ''))
        else:
            _texto(fila.cells[1], '')
        doc.add_paragraph()

    bloque_transversalidad(doc, corte, voc)


def bloque_transversalidad(doc, corte, voc):
    tr = corte.get('transversalidad', {}) or {}
    t = _tabla(doc, 1, 4, anchos=[6.0, 3.4, 3.1, 3.1])
    _fusionar_fila(t, 0, 'Transversalidad')

    fila = t.add_row()
    c = fila.cells[0]
    for otra in fila.cells[1:]:
        c = c.merge(otra)
    _sub(c, 'Currículum Fundamental')

    fila = t.add_row()
    _sub(fila.cells[0], voc['transversal_fundamental'])
    _sub(fila.cells[1].merge(fila.cells[3]), voc['transversal_vinculo'])

    for v in (tr.get('fundamental', []) or [{}]):
        fila = t.add_row()
        _texto(fila.cells[0], v.get('asignatura', ''))
        _texto(fila.cells[1].merge(fila.cells[3]), v.get('vinculo', ''))

    fila = t.add_row()
    _sub(fila.cells[0].merge(fila.cells[1]), 'Habilidades para la Vida y el Trabajo')
    _sub(fila.cells[2].merge(fila.cells[3]), voc['sostenible'])

    fila = t.add_row()
    _texto(fila.cells[0].merge(fila.cells[1]), tr.get('hvyt', ''))
    _texto(fila.cells[2].merge(fila.cells[3]), tr.get('sostenible', ''))
    doc.add_paragraph()


def bloque_fase(doc, fase, voc, etiqueta_fase, etiqueta_eval):
    """Una tabla por fase: actividades, recursos y evaluacion."""
    t = _tabla(doc, 1, 4, anchos=[6.0, 3.4, 3.3, 3.3])
    _fusionar_fila(t, 0, 'Actividades de aprendizaje')

    fila = t.add_row()
    a = fila.cells[0].merge(fila.cells[2])
    _sub(a, '%s (describa las actividades)' % etiqueta_fase)
    _sub(fila.cells[3], 'Tiempo')

    for act in fase.get('actividades', []) or [{}]:
        fila = t.add_row()
        _texto(fila.cells[0].merge(fila.cells[2]), act.get('texto', ''))
        _texto(fila.cells[3], act.get('tiempo', ''), centrado=True)

    fila = t.add_row()
    c = fila.cells[0]
    for otra in fila.cells[1:]:
        c = c.merge(otra)
    _sub(c, 'Recursos didácticos y tecnológicos')

    fila = t.add_row()
    c = fila.cells[0]
    for otra in fila.cells[1:]:
        c = c.merge(otra)
    _texto(c, fase.get('recursos', ''))

    fila = t.add_row()
    c = fila.cells[0]
    for otra in fila.cells[1:]:
        c = c.merge(otra)
    _enc(c, 'Evaluación')

    ev = fase.get('evaluacion', {}) or {}
    fila = t.add_row()
    a = fila.cells[0].merge(fila.cells[2])
    _sub(a, etiqueta_eval)
    _sub(fila.cells[3], 'Tiempo')
    fila = t.add_row()
    _texto(fila.cells[0].merge(fila.cells[2]), ev.get('descripcion', ''))
    _texto(fila.cells[3], ev.get('tiempo', ''), centrado=True)

    fila = t.add_row()
    for i, nombre in enumerate(
            ('Evidencia', 'Criterios', 'Instrumentos', 'Ponderación')):
        _sub(fila.cells[i], nombre)

    for r in ev.get('filas', []) or []:
        fila = t.add_row()
        for i in range(4):
            _texto(fila.cells[i], r[i] if i < len(r) else '',
                   centrado=(i == 3))

    # Retroalimentacion, solo en el cierre
    if fase.get('retroalimentacion'):
        rt = fase['retroalimentacion']
        fila = t.add_row()
        _sub(fila.cells[0], 'Retroalimentación')
        _sub(fila.cells[1].merge(fila.cells[2]), 'Actividades')
        _sub(fila.cells[3], 'Tiempo')
        fila = t.add_row()
        _texto(fila.cells[0], rt.get('titulo', ''))
        _texto(fila.cells[1].merge(fila.cells[2]), rt.get('actividades', ''))
        _texto(fila.cells[3], rt.get('tiempo', ''), centrado=True)

    doc.add_paragraph()


# --------------------------------------------------------------------------
# Documento completo
# --------------------------------------------------------------------------

def generar(spec, salida):
    voc = T.obtener(spec.get('modelo', 'uac'))
    ident = spec.get('identificacion', {})

    d = dc.Doc(perfil=PERFIL_NOMBRE, pie=spec.get('pie', 'Colegio de Bachilleres'))
    doc = d.d

    # Autoria del archivo: el nombre de quien imparte la asignatura, no el de la
    # herramienta. Es el dato que Word muestra en propiedades del documento.
    docente = (ident.get('docente') or '').strip()
    props = doc.core_properties
    if docente:
        props.author = docente
        props.last_modified_by = docente
    props.title = 'Planeación didáctica · %s' % (ident.get('unidad') or '')
    props.subject = 'Colegio de Bachilleres · Componente de Formación Laboral'
    props.category = voc['etiqueta_modelo']

    tit = doc.add_paragraph()
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tit.add_run('Planeación didáctica')
    r.font.name = FUENTE
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = dc._rgb(VERDE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('%s · %s' % (ident.get('unidad', ''), voc['etiqueta_modelo']))
    r.font.name = FUENTE
    r.font.size = Pt(10)

    bloque_identificacion(doc, ident, voc)

    for corte in spec.get('cortes', []):
        p = doc.add_paragraph()
        r = p.add_run('CORTE %s' % corte.get('numero', ''))
        r.font.name = FUENTE
        r.font.size = Pt(12)
        r.font.bold = True

        bloque_curricular(doc, corte, voc)
        bloque_desarrollos(doc, corte, voc)

        for clave, etq, ev in (
            ('apertura', voc['fase_apertura'], 'Diagnóstica'),
            ('desarrollo', voc['fase_desarrollo'], 'Formativa'),
            ('cierre', voc['fase_cierre'], 'Sumativa'),
        ):
            if corte.get(clave):
                bloque_fase(doc, corte[clave], voc, etq, ev)

    d.guardar(salida)
    return salida


ESQUEMA = {
    "modelo": "uac",
    "pie": "Colegio de Bachilleres",
    "identificacion": {
        "institucion": "Colegio de Bachilleres",
        "docente": "", "plantel": "", "curriculum": "Laboral",
        "tob": "", "unidad": "", "semestre": "", "ciclo": "",
        "horas_semestre": "", "fecha": ""
    },
    "cortes": [{
        "numero": 1, "nombre_corte": "", "horas_corte": "",
        "competencia_laboral": "", "proposito": "", "meta_corte": "",
        "desarrollos": [
            {"texto": "", "conocimiento": "", "desempeno": "", "producto": ""}
        ],
        "transversalidad": {
            "fundamental": [{"asignatura": "", "vinculo": ""}],
            "hvyt": "", "sostenible": ""
        },
        "apertura": {
            "actividades": [{"texto": "", "tiempo": ""}],
            "recursos": "",
            "evaluacion": {"descripcion": "", "tiempo": "",
                           "filas": [["", "", "", "0%"]]}
        },
        "desarrollo": {
            "actividades": [{"texto": "", "tiempo": ""}],
            "recursos": "",
            "evaluacion": {"descripcion": "", "tiempo": "",
                           "filas": [["", "", "", "0%"]]}
        },
        "cierre": {
            "actividades": [{"texto": "", "tiempo": ""}],
            "recursos": "",
            "evaluacion": {"descripcion": "", "tiempo": "",
                           "filas": [["", "", "", "25%"]]},
            "retroalimentacion": {"titulo": "", "actividades": "", "tiempo": ""}
        }
    }]
}


def main():
    if '--esquema' in sys.argv:
        sys.stdout.write(json.dumps(ESQUEMA, ensure_ascii=False, indent=2))
        sys.stdout.write('\n')
        return
    if len(sys.argv) < 3:
        sys.stderr.write(__doc__)
        sys.exit(2)
    with open(sys.argv[1], 'rb') as f:
        spec = json.loads(f.read().decode('utf-8'))
    ruta = generar(spec, sys.argv[2])
    sys.stdout.write('Planeacion generada: %s\n' % ruta)


if __name__ == '__main__':
    main()

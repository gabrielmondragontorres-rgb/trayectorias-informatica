# -*- coding: utf-8 -*-
"""
docx-cobach — Motor de documentos Word con identidad del Colegio de Bachilleres.

Uso minimo:

    from docx_cobach import Doc

    d = Doc(perfil="guia", pie="Guia de estudio - Colegio de Bachilleres")
    d.portada("Guia de estudio", "Aplicacion del comercio electronico",
              "Quinto semestre", [("Plantel", "16 Tlahuac")])
    d.h1("Presentacion")
    d.p("Texto del parrafo, con **negritas** en linea.")
    d.tabla(["Meta", "Evidencia"], [["Desarrollar...", "Landing page"]])
    d.guardar("salida.docx")

Este motor SOLO dibuja. Que secciones lleva un documento lo mandan las skills
`estructura-guia` y `formato-word-guia`.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ===========================================================================
# APRENDIZAJE BLINDADO — orden de elementos en OOXML
#
# El esquema fija el orden de los hijos de <w:pPr> y <w:tcPr>. Insertar
# <w:shd> o <w:pBdr> al final produce un archivo que Word abre con el aviso
# "documento danado" y ofrece repararlo. Ambos deben ir ANTES de w:spacing,
# w:ind y w:jc. En w:tcPr, w:shd va despues de w:tcBorders.
#
# Detectado el 19 de agosto de 2026: seis documentos ya producidos habrian
# salido corruptos. Nunca usar pPr.append(shd); usar _colocar().
# ===========================================================================
_PPR_SEQ = (
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
    "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap", "w:overflowPunct",
    "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN", "w:bidi", "w:adjustRightInd",
    "w:snapToGrid", "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
    "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
    "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle", "w:rPr",
    "w:sectPr", "w:pPrChange",
)
_TCPR_SEQ = (
    "w:cnfStyle", "w:tcW", "w:gridSpan", "w:hMerge", "w:vMerge", "w:tcBorders",
    "w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign",
    "w:hideMark", "w:headers", "w:cellIns", "w:cellDel", "w:cellMerge", "w:tcPrChange",
)


def _colocar(padre, elemento, tag, secuencia):
    """Inserta el elemento donde el esquema lo exige, sin duplicarlo."""
    previo = padre.find(qn(tag))
    if previo is not None:
        padre.remove(previo)
    padre.insert_element_before(elemento, *secuencia[secuencia.index(tag) + 1:])


def _shd(hexcolor):
    e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear")
    e.set(qn("w:color"), "auto")
    e.set(qn("w:fill"), hexcolor)
    return e


def sombrear_celda(celda, hexcolor):
    _colocar(celda._tc.get_or_add_tcPr(), _shd(hexcolor), "w:shd", _TCPR_SEQ)


def sombrear_parrafo(p, hexcolor):
    _colocar(p._p.get_or_add_pPr(), _shd(hexcolor), "w:shd", _PPR_SEQ)


def borde_parrafo(p, color, lado="left", sz=18, space=10):
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:" + lado)
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), str(space))
    b.set(qn("w:color"), color)
    pbdr.append(b)
    _colocar(p._p.get_or_add_pPr(), pbdr, "w:pBdr", _PPR_SEQ)


def borde_tabla_negro(tabla, sz=4):
    """Contorno negro fino y sin relleno: el cuadro 1x1 de metas especificas."""
    tbl_pr = tabla._tbl.tblPr
    bordes = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement("w:" + lado)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), "000000")
        bordes.append(e)
    tbl_pr.append(bordes)


# ===========================================================================
# PERFILES VISUALES
#
# Decision del Mtro. Luis Gabriel el 19 de agosto de 2026: el perfil `guia`
# usa el MISMO verde institucional que la plataforma. Antes de esa fecha la
# especificacion de guias marcaba azul #2F5496, que es el acento por defecto
# de Word, no un color de marca del Colegio. El verde #006837 si es el del
# logotipo.
#
# Lo que el perfil `guia` conserva de `formato-word-guia`: los tamanos
# (14 pt H1, 12 pt H2), el interlineado 1.15 con 2 pt posterior, la
# tipografia de 10 pt dentro de tablas con interlineado sencillo, y el
# cebreado verde menta #D1EBE5 (que ya era verde y combina).
# ===========================================================================
PERFILES = {
    "documento-tecnico": {
        "fuente": "Noto Sans",
        "cuerpo_pt": 11,
        "interlineado": 1.15,
        "espacio_posterior_pt": 6,
        "h1_pt": 16, "h1_color": "FFFFFF", "h1_fondo": "006837",
        "h2_pt": 13.5, "h2_color": "006837", "h2_regla": "00A859",
        "h3_pt": 11.5, "h3_color": "007A42",
        "tabla_encabezado_fondo": "006837",
        "tabla_cebra": "F4F9F6",
        "tabla_pt": 9.5,
        "ficha_etiqueta_fondo": "F1F8F4",
        "acento": "00A859",
        "color_enlace": "0563C1",
        "margenes_cm": (2.2, 2.2, 2.4, 2.4),
    },
    "guia": {
        "fuente": "Noto Sans",
        "cuerpo_pt": 11,
        "interlineado": 1.15,
        "espacio_posterior_pt": 2,
        "h1_pt": 14, "h1_color": "006837", "h1_fondo": None,
        "h2_pt": 12, "h2_color": "006837", "h2_regla": None,
        "h3_pt": 11, "h3_color": "007A42",
        "tabla_encabezado_fondo": "006837",
        "tabla_cebra": "D1EBE5",
        "tabla_pt": 10,
        "ficha_etiqueta_fondo": "EAF5F0",
        "acento": "00A859",
        "color_enlace": "0563C1",
        "margenes_cm": (3.75, 2.54, 1.91, 1.91),
    },
}

BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
GRIS = RGBColor(0x44, 0x44, 0x44)
GRIS_CLARO = RGBColor(0x66, 0x66, 0x66)


def _rgb(hexcolor):
    return RGBColor(int(hexcolor[0:2], 16), int(hexcolor[2:4], 16), int(hexcolor[4:6], 16))


# ------------------------------------------------------- verificacion de color
def _luminancia(hexcolor):
    c = hexcolor.lstrip("#")
    v = []
    for i in (0, 2, 4):
        x = int(c[i:i + 2], 16) / 255
        v.append(x / 12.92 if x <= 0.03928 else ((x + 0.055) / 1.055) ** 2.4)
    return 0.2126 * v[0] + 0.7152 * v[1] + 0.0722 * v[2]


def contraste(a, b):
    """Ratio de contraste WCAG entre dos colores hexadecimales."""
    l1, l2 = sorted((_luminancia(a), _luminancia(b)), reverse=True)
    return (l1 + 0.05) / (l2 + 0.05)


def verificar_perfil(nombre):
    """Comprueba que las combinaciones de color del perfil alcancen 4.5:1."""
    p = PERFILES[nombre]
    pares = [
        ("Texto de encabezado de tabla sobre su fondo", "FFFFFF", p["tabla_encabezado_fondo"]),
        ("Texto negro sobre el cebreado de filas", "1A1A1A", p["tabla_cebra"]),
        ("Titulo 2 sobre blanco", p["h2_color"], "FFFFFF"),
        ("Titulo 3 sobre blanco", p["h3_color"], "FFFFFF"),
        ("Texto negro sobre el fondo de etiqueta de ficha", "1A1A1A", p["ficha_etiqueta_fondo"]),
    ]
    if p["h1_fondo"]:
        pares.insert(0, ("Titulo 1 sobre su banda de color", p["h1_color"], p["h1_fondo"]))
    else:
        pares.insert(0, ("Titulo 1 sobre blanco", p["h1_color"], "FFFFFF"))
    filas = []
    for etq, fg, bg in pares:
        r = contraste(fg, bg)
        filas.append((etq, fg, bg, r, r >= 4.5))
    return filas


# ===========================================================================
class Doc:
    """Documento Word con identidad COBACH. Solo dibuja; no define estructura."""

    def __init__(self, perfil="documento-tecnico", pie=""):
        if perfil not in PERFILES:
            raise ValueError(f"Perfil desconocido: {perfil}. Use {list(PERFILES)}")
        self.perfil = perfil
        self.cfg = PERFILES[perfil]
        self.d = Document()
        self._n_figura = 0
        self._estilos()
        if pie:
            self._pie(pie)

    # ---------------------------------------------------------------- base
    def _estilos(self):
        c = self.cfg
        n = self.d.styles["Normal"]
        n.font.name = c["fuente"]
        n.font.size = Pt(c["cuerpo_pt"])
        n.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        # Sin esto, Word Online ignora la fuente y sustituye por Calibri.
        n.element.rPr.rFonts.set(qn("w:eastAsia"), c["fuente"])
        n.paragraph_format.space_after = Pt(c["espacio_posterior_pt"])
        n.paragraph_format.line_spacing = c["interlineado"]
        sup, inf, izq, der = c["margenes_cm"]
        for s in self.d.sections:
            s.top_margin, s.bottom_margin = Cm(sup), Cm(inf)
            s.left_margin, s.right_margin = Cm(izq), Cm(der)

    def _pie(self, texto):
        for s in self.d.sections:
            p = s.footer.paragraphs[0]
            p.text = texto
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(7.5)
                r.font.color.rgb = GRIS_CLARO
                r.font.name = self.cfg["fuente"]

    def _runs(self, par, texto, tam=None, italica=False, negrita=False, color=None):
        """Escribe el texto interpretando **negritas** en linea."""
        tam = tam or self.cfg["cuerpo_pt"]
        for i, parte in enumerate(str(texto).split("**")):
            if parte == "":
                continue
            r = par.add_run(parte)
            r.font.size = Pt(tam)
            r.font.name = self.cfg["fuente"]
            r.italic = italica
            r.bold = negrita or (i % 2 == 1)
            if color is not None:
                r.font.color.rgb = color
            r._element.get_or_add_rPr().get_or_add_rFonts().set(
                qn("w:eastAsia"), self.cfg["fuente"])
        return par

    # ------------------------------------------------------------- bloques
    def portada(self, titulo, subtitulo="", sobretitulo="", metadatos=None):
        c = self.cfg
        if sobretitulo:
            p = self.d.add_paragraph()
            p.paragraph_format.space_before = Pt(80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(sobretitulo.upper())
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.color.rgb = _rgb(c["acento"])
            r.font.name = c["fuente"]

        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10 if sobretitulo else 80)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(titulo)
        r.bold = True
        r.font.size = Pt(25)
        r.font.color.rgb = _rgb(c["h1_color"] if c["h1_fondo"] is None else "006837")
        r.font.name = c["fuente"]

        if subtitulo:
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(26)
            r = p.add_run(subtitulo)
            r.font.size = Pt(13)
            r.font.color.rgb = GRIS
            r.font.name = c["fuente"]

        if metadatos:
            self.ficha(metadatos, ancho_etq=4.6, ancho_val=11.0, tam=9.5)
        self.salto()

    def h1(self, texto, salto_antes=True):
        c = self.cfg
        if salto_antes:
            self.salto()
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(10)
        if c["h1_fondo"]:
            sombrear_parrafo(p, c["h1_fondo"])
            texto = "  " + texto
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(c["h1_pt"])
        r.font.color.rgb = _rgb(c["h1_color"])
        r.font.name = c["fuente"]
        return p

    def h2(self, texto):
        c = self.cfg
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(c["h2_pt"])
        r.font.color.rgb = _rgb(c["h2_color"])
        r.font.name = c["fuente"]
        if c["h2_regla"]:
            borde_parrafo(p, c["h2_regla"], "bottom", sz=8, space=4)
        return p

    def h3(self, texto):
        c = self.cfg
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(texto)
        r.bold = True
        r.font.size = Pt(c["h3_pt"])
        r.font.color.rgb = _rgb(c["h3_color"])
        r.font.name = c["fuente"]
        return p

    def p(self, texto, tam=None, italica=False, negrita=False, color=None,
          espacio=None, justificado=True):
        par = self.d.add_paragraph()
        par.paragraph_format.space_after = Pt(
            self.cfg["espacio_posterior_pt"] if espacio is None else espacio)
        if justificado:
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return self._runs(par, texto, tam, italica, negrita, color)

    def vineta(self, texto, tam=None, nivel=0):
        par = self.d.add_paragraph(style="List Bullet")
        par.paragraph_format.space_after = Pt(3)
        par.paragraph_format.left_indent = Cm(0.75 + 0.6 * nivel)
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return self._runs(par, texto, tam)

    def numerada(self, texto, tam=None):
        par = self.d.add_paragraph(style="List Number")
        par.paragraph_format.space_after = Pt(3)
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return self._runs(par, texto, tam)

    def etiqueta(self, texto, fondo=None):
        """Franja de color con texto blanco: rotula el bloque que sigue."""
        p = self.d.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(0)
        sombrear_parrafo(p, fondo or self.cfg["acento"])
        r = p.add_run("  " + texto.upper())
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = BLANCO
        r.font.name = self.cfg["fuente"]
        return p

    def cita(self, texto, tam=None):
        p = self.d.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        borde_parrafo(p, self.cfg["acento"], "left", sz=18, space=10)
        return self._runs(p, texto, tam, italica=True)

    def aviso(self, titulo, texto, fondo="FFF7E6", borde="E0A800"):
        for i, (t, negrita) in enumerate([(titulo, True), (texto, False)]):
            p = self.d.add_paragraph()
            p.paragraph_format.space_before = Pt(10 if i == 0 else 0)
            p.paragraph_format.space_after = Pt(2 if i == 0 else 8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            sombrear_parrafo(p, fondo)
            borde_parrafo(p, borde, "left", sz=18, space=8)
            self._runs(p, "  " + t, 10, negrita=negrita)

    def bloque_prompt(self, texto):
        """Texto literal sobre fondo oscuro: prompts que no deben editarse."""
        self._monoespaciado(texto, "1E293B", RGBColor(0x8B, 0xEB, 0xBA), 8.5, Cm(0.25))

    def bloque_codigo(self, texto):
        self._monoespaciado(texto, "F3F5F7", RGBColor(0x1F, 0x2D, 0x3D), 7.5, Cm(0.25))

    def _monoespaciado(self, texto, fondo, color, tam, sangria):
        for linea in str(texto).split("\n"):
            p = self.d.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = sangria
            p.paragraph_format.right_indent = sangria
            sombrear_parrafo(p, fondo)
            r = p.add_run(linea if linea.strip() else " ")
            r.font.size = Pt(tam)
            r.font.name = "Consolas"
            r.font.color.rgb = color
        self.espacio(4)

    def tabla(self, encabezados, filas, anchos=None, tam=None, cebra=True):
        c = self.cfg
        tam = tam or c["tabla_pt"]
        t = self.d.add_table(rows=1, cols=len(encabezados))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(encabezados):
            cel = t.rows[0].cells[i]
            cel.text = ""
            r = cel.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(tam)
            r.font.color.rgb = BLANCO
            r.font.name = c["fuente"]
            sombrear_celda(cel, c["tabla_encabezado_fondo"])
        for idx, fila in enumerate(filas):
            celdas = t.add_row().cells
            for i, v in enumerate(fila):
                celdas[i].text = ""
                p = celdas[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                self._runs(p, str(v), tam)
                if cebra and idx % 2 == 1:
                    sombrear_celda(celdas[i], c["tabla_cebra"])
        if anchos:
            for row in t.rows:
                for i, w in enumerate(anchos):
                    row.cells[i].width = Cm(w)
        self.espacio(4)
        return t

    def ficha(self, campos, ancho_etq=5.0, ancho_val=10.6, tam=None):
        """Tabla de dos columnas etiqueta / valor."""
        c = self.cfg
        tam = tam or c["tabla_pt"]
        t = self.d.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for k, v in campos:
            celdas = t.add_row().cells
            celdas[0].text = ""
            celdas[1].text = ""
            r = celdas[0].paragraphs[0].add_run(k)
            r.bold = True
            r.font.size = Pt(tam)
            r.font.color.rgb = _rgb(c["h2_color"])
            r.font.name = c["fuente"]
            sombrear_celda(celdas[0], c["ficha_etiqueta_fondo"])
            p = celdas[1].paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._runs(p, str(v), tam)
        for row in t.rows:
            row.cells[0].width = Cm(ancho_etq)
            row.cells[1].width = Cm(ancho_val)
        self.espacio(4)
        return t

    def cuadro_meta(self, texto, prefijo=""):
        """Tabla 1x1 de metas especificas: contorno negro, sin relleno.

        Regla de `formato-word-guia`: todo el texto sin negrita salvo el
        prefijo literal "Meta especifica N".
        """
        t = self.d.add_table(rows=1, cols=1)
        borde_tabla_negro(t)
        cel = t.rows[0].cells[0]
        cel.text = ""
        p = cel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        if prefijo:
            r = p.add_run(prefijo + " ")
            r.bold = True
            r.font.size = Pt(self.cfg["tabla_pt"])
            r.font.name = self.cfg["fuente"]
        self._runs(p, texto, self.cfg["tabla_pt"])
        self.espacio(4)
        return t

    def imagen(self, ruta, ancho_cm=15.5, pie="", numerar=True):
        self.d.add_picture(ruta, width=Cm(ancho_cm))
        self.d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if pie:
            self._n_figura += 1
            p = self.d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            texto = f"Figura {self._n_figura}. {pie}" if numerar else pie
            r = p.add_run(texto)
            r.font.size = Pt(8.5)
            r.italic = True
            r.font.color.rgb = GRIS_CLARO
            r.font.name = self.cfg["fuente"]

    def espacio(self, pt=8):
        self.d.add_paragraph().paragraph_format.space_after = Pt(pt)

    def salto(self):
        self.d.add_page_break()

    def guardar(self, ruta):
        carpeta = os.path.dirname(os.path.abspath(ruta))
        os.makedirs(carpeta, exist_ok=True)
        self.d.save(ruta)
        return ruta

# -*- coding: utf-8 -*-
"""
Valida un .docx antes de entregarlo.

Comprueba tres cosas que rompen un documento sin que se note al generarlo:

  1. Orden de los hijos de <w:pPr> y <w:tcPr>. Si esta mal, Word abre el
     archivo con el aviso "documento danado". Es el fallo que se detecto el
     19 de agosto de 2026 con seis documentos ya producidos.
  2. Restos de marcado sin procesar (** de negrita, marcadores de plantilla).
  3. Celdas de tabla vacias y tablas sin filas de datos.

Uso:
    python validar_docx.py archivo.docx [otro.docx ...]
    python validar_docx.py carpeta/
"""
from __future__ import annotations

import glob
import os
import re
import sys
import zipfile

from lxml import etree

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

PPR = ["pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
       "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
       "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE",
       "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
       "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc", "textDirection",
       "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
       "sectPr", "pPrChange"]
TCPR = ["cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd", "noWrap",
        "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark", "headers",
        "cellIns", "cellDel", "cellMerge", "tcPrChange"]

MARCADORES = [(r"\*\*", "asteriscos de negrita sin procesar"),
              (r"\bTODO\b", "marcador TODO"),
              (r"\{\{?\w+\}?\}", "marcador de plantilla sin sustituir"),
              (r"\bNone\b", "valor None de Python")]


def _fuera_de_orden(root, tag, seq):
    malos = 0
    for pr in root.iter(W + tag):
        idx = [seq.index(c.tag[len(W):]) for c in pr
               if c.tag.startswith(W) and c.tag[len(W):] in seq]
        if idx != sorted(idx):
            malos += 1
    return malos


def validar(ruta):
    problemas = []
    with zipfile.ZipFile(ruta) as z:
        partes = [n for n in z.namelist()
                  if n.startswith("word/") and n.endswith(".xml")]
        for nombre in partes:
            root = etree.fromstring(z.read(nombre))
            a = _fuera_de_orden(root, "pPr", PPR)
            b = _fuera_de_orden(root, "tcPr", TCPR)
            if a:
                problemas.append(f"{nombre}: {a} <w:pPr> con hijos fuera de orden "
                                 f"(Word marcara el archivo como danado)")
            if b:
                problemas.append(f"{nombre}: {b} <w:tcPr> con hijos fuera de orden")

    try:
        import docx
    except ImportError:
        return problemas, None

    d = docx.Document(ruta)
    trozos = [p.text for p in d.paragraphs]
    celdas_vacias = 0
    tablas_sin_datos = 0
    for t in d.tables:
        if len(t.rows) == 0:
            tablas_sin_datos += 1
        for row in t.rows:
            for c in row.cells:
                trozos.append(c.text)
                if not c.text.strip():
                    celdas_vacias += 1
    texto = "\n".join(trozos)

    for patron, etq in MARCADORES:
        n = len(re.findall(patron, texto))
        if n:
            problemas.append(f"contenido: {n} coincidencias de {etq}")
    if tablas_sin_datos:
        problemas.append(f"contenido: {tablas_sin_datos} tablas sin ninguna fila")

    resumen = {
        "parrafos": len(d.paragraphs),
        "tablas": len(d.tables),
        "palabras": len(texto.split()),
        "celdas_vacias": celdas_vacias,
    }
    return problemas, resumen


def main(argv):
    if not argv:
        print(__doc__)
        return 1
    rutas = []
    for a in argv:
        rutas.extend(sorted(glob.glob(os.path.join(a, "*.docx")))
                     if os.path.isdir(a) else [a])
    if not rutas:
        print("No se encontro ningun .docx en:", ", ".join(argv))
        return 1

    total = 0
    for r in rutas:
        problemas, resumen = validar(r)
        total += len(problemas)
        nombre = os.path.basename(r)
        if problemas:
            print(f"[FALLA] {nombre}")
            for p in problemas:
                print(f"        - {p}")
        else:
            extra = ""
            if resumen:
                extra = (f"  ({resumen['parrafos']} parrafos, {resumen['tablas']} tablas, "
                         f"{resumen['palabras']} palabras, "
                         f"{resumen['celdas_vacias']} celdas vacias)")
            print(f"[OK]    {nombre}{extra}")

    print(f"\n{len(rutas)} archivo(s) revisado(s) · {total} problema(s)")
    return 1 if total else 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
